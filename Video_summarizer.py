import os
from moviepy.editor import VideoFileClip
import speech_recognition as sr
import openai
import time
from docx import Document


class VideoSummarizer:
    def __init__(self, api_key):
        self.recognizer = sr.Recognizer()
        openai.api_key = api_key

    def extract_audio(self, video_path, audio_output="temp_audio.wav"):
        print("Extracting audio from video...")
        video = VideoFileClip(video_path)
        video.audio.write_audiofile(audio_output, codec='pcm_s16le')
        return audio_output

    def transcribe_chunk_with_retry(self, audio_chunk, max_retries=3, delay=2):
        """Transcribe audio chunk with retry mechanism."""
        for attempt in range(max_retries):
            try:
                text = self.recognizer.recognize_google(audio_chunk)
                return text
            except sr.RequestError as e:
                if attempt == max_retries - 1:
                    print(f"Final attempt failed: {str(e)}")
                    return ""
                print(f"Attempt {attempt + 1} failed, retrying in {delay} seconds...")
                time.sleep(delay)
                delay *= 2
            except sr.UnknownValueError:
                return ""

    def transcribe_audio(self, audio_path, chunk_duration=30):
        print("Transcribing audio...")
        audio = sr.AudioFile(audio_path)
        full_transcript = []

        with audio as source:
            duration = source.DURATION
            chunks = int(duration / chunk_duration) + 1

            for i in range(chunks):
                try:
                    print(f"Processing chunk {i + 1}/{chunks}...")
                    audio_chunk = self.recognizer.record(source, duration=chunk_duration)
                    text = self.transcribe_chunk_with_retry(audio_chunk)
                    if text:
                        full_transcript.append(text)
                except Exception as e:
                    print(f"Error in chunk {i}: {str(e)}")
                    continue
                time.sleep(0.5)

        return " ".join(full_transcript)

    def generate_structured_summary(self, transcript, max_tokens=2000):
        print("Generating structured summary...")
        try:
            system_prompt = """You are a professional note-taker and summarizer. Create a detailed, structured summary of the content with the following format:

# [Title of the Content]

## Overview
[2-3 sentences summarizing the main topic and key points]

## Key Points
- [Main point 1]
- [Main point 2]
- [Main point 3]

## Detailed Notes

### [Main Topic 1]
#### [Subtopic 1.1]
- [Detailed point]
- [Detailed point]
- [Supporting information]

#### [Subtopic 1.2]
- [Detailed point]
- [Detailed point]
- [Supporting information]

### [Main Topic 2]
#### [Subtopic 2.1]
- [Detailed point]
- [Detailed point]
- [Supporting information]

## Summary and Conclusions
- [Key takeaway 1]
- [Key takeaway 2]
- [Key takeaway 3]"""

            response = openai.ChatCompletion.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": transcript}
                ],
                max_tokens=max_tokens,
                temperature=0.7
            )
            return response.choices[0].message['content']
        except Exception as e:
            return f"Error generating summary: {str(e)}"

    def save_summary_as_markdown(self, summary, output_path):
        with open(f"{output_path}.md", 'w', encoding='utf-8') as f:
            f.write(summary)

    def save_summary_as_word(self, summary, output_path):
        doc = Document()
        for line in summary.split('\n'):
            if line.startswith('#'):
                if line.startswith('###'):
                    doc.add_heading(line.strip('#').strip(), level=3)
                elif line.startswith('##'):
                    doc.add_heading(line.strip('#').strip(), level=2)
                else:
                    doc.add_heading(line.strip('#').strip(), level=1)
            else:
                doc.add_paragraph(line)
        doc.save(f"{output_path}.docx")

    def process_video(self, video_path, output_format, output_path="content_summary"):
        try:
            # Extract audio
            audio_path = self.extract_audio(video_path)

            transcript = self.transcribe_audio(audio_path)

            summary = self.generate_structured_summary(transcript)

            # Save output in the selected format
            if output_format == "markdown":
                self.save_summary_as_markdown(summary, output_path)
                print(f"Summary saved as {output_path}.md")
            elif output_format == "word":
                self.save_summary_as_word(summary, output_path)
                print(f"Summary saved as {output_path}.docx")

            os.remove(audio_path)

            return summary

        except Exception as e:
            print(f"Error processing video: {str(e)}")
            return None

    @staticmethod
    def get_video_duration(video_path):
        video = VideoFileClip(video_path)
        return video.duration  # Duration in seconds

    